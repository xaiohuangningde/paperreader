from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from io import BytesIO
import matplotlib.pyplot as plt
import matplotlib as mpl
from matplotlib import rc
import re

class WordReportGenerator:
    def __init__(self):
        self.document = Document()
        # 设置页面边距，方便放宽表格
        section = self.document.sections[0]
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        
        # 添加标题
        heading = self.document.add_heading('SPE 文献深度分析报告', 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加生成时间和说明
        time_para = self.document.add_paragraph()
        time_run = time_para.add_run(f"生成时间: {self._get_current_time()}")
        time_run.font.size = Pt(10)
        time_run.font.color.rgb = RGBColor(128, 128, 128)
        time_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加说明
        note_para = self.document.add_paragraph()
        note_run = note_para.add_run("注: 本报告由 DeepSpec Pro 自动生成，结论与参数均源自原文。")
        note_run.font.size = Pt(9)
        note_run.font.color.rgb = RGBColor(128, 128, 128)
        note_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加空行
        self.document.add_paragraph()

    def _get_current_time(self):
        """获取当前时间字符串"""
        from datetime import datetime
        return datetime.now().strftime("%Y年%m月%d日 %H:%M")
    
    def _set_cell_background(self, cell, color):
        """设置单元格背景颜色"""
        shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color))
        cell._tc.get_or_add_tcPr().append(shading_elm)
    
    def _render_latex_to_image(self, latex_string, fontsize=12):
        """
        将LaTeX公式渲染为图片
        """
        try:
            # 配置matplotlib以支持中文和LaTeX
            plt.rcParams['font.sans-serif'] = ['SimHei', 'DejaVu Sans']
            plt.rcParams['axes.unicode_minus'] = False
            
            # 创建一个简单的图形
            fig, ax = plt.subplots(figsize=(6, 1))
            ax.axis('off')  # 隐藏坐标轴
            
            # 提取LaTeX字符串（去掉$符号）
            clean_latex = latex_string.replace('$$', '').replace('$', '').strip()
            
            # 如果包含数学符号，尝试渲染
            if '\\' in clean_latex or '_' in clean_latex or '^' in clean_latex:
                try:
                    # 尝试渲染LaTeX
                    ax.text(0.5, 0.5, f'${clean_latex}$', 
                           fontsize=fontsize, ha='center', va='center')
                except:
                    # 如果LaTeX渲染失败，直接显示文本
                    ax.text(0.5, 0.5, clean_latex, 
                           fontsize=fontsize, ha='center', va='center')
            else:
                # 如果是简单公式，直接显示
                ax.text(0.5, 0.5, clean_latex, 
                       fontsize=fontsize, ha='center', va='center')
            
            # 将图形保存到字节流
            img_buffer = BytesIO()
            plt.savefig(img_buffer, format='png', bbox_inches='tight', dpi=150)
            img_buffer.seek(0)
            plt.close()
            
            return img_buffer
        except Exception as e:
            print(f"LaTeX渲染失败: {str(e)}")
            return None

    def add_paper_analysis(self, data, image_stream=None):
        """
        添加论文分析到报告
        
        Args:
            data: 字典，包含 'title', 'purpose', 'conclusion', 'params', 'formulas', 'comments', 'why', 'page_source'
            image_stream: 图片的字节流 (BytesIO)，如果没有则为 None
        """
        # 添加一个分隔符
        self.document.add_paragraph("-" * 50)
        
        # 创建表格：1行5列 (对应要求的docx结构)
        table = self.document.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        table.autofit = False 
        
        # 设置表头
        hdr_cells = table.rows[0].cells
        headers = ["Article", "具体内容 (1): 目的与结论", "具体内容 (2): 参数与公式", "Comments (想法)", "Why"]
        
        for i, text in enumerate(headers):
            run = hdr_cells[i].paragraphs[0].add_run(text)
            run.bold = True
            run.font.size = Pt(10)
            # 给表头加背景色 (灰色)
            self._set_cell_background(hdr_cells[i], "E7E6E6")
        
        # 填入数据行
        row_cells = table.add_row().cells
        
        # Col 1: Article Title
        title_para = row_cells[0].paragraphs[0]
        title_run = title_para.add_run(data.get('title', 'N/A'))
        title_run.bold = True
        title_run.font.size = Pt(9)
        
        # Col 2: 目的与结论 (Specifics 1)
        p1 = row_cells[1].paragraphs[0]
        
        # 添加页码来源
        if 'page_source' in data and data['page_source']:
            source_run = p1.add_run(f"[来源: 第{data['page_source']}页]\n")
            source_run.font.size = Pt(8)
            source_run.font.color.rgb = RGBColor(100, 100, 100)
        
        # 添加目的
        purpose_run = p1.add_run("\n1. 研究目的：\n")
        purpose_run.bold = True
        purpose_run.font.size = Pt(9)
        purpose_content = p1.add_run(f"{data.get('purpose', 'N/A')}")
        purpose_content.font.size = Pt(9)
        
        # 添加结论
        conclusion_run = p1.add_run("\n\n2. 核心结论：\n")
        conclusion_run.bold = True
        conclusion_run.font.size = Pt(9)
        
        conclusions = data.get('conclusion', [])
        if isinstance(conclusions, str):
            conclusions = [conclusions]
            
        for idx, point in enumerate(conclusions, 1):
            conclusion_num = p1.add_run(f"({idx}) ")
            conclusion_num.bold = True
            conclusion_num.font.size = Pt(9)
            conclusion_content = p1.add_run(f"{point}\n")
            conclusion_content.font.size = Pt(9)
            
        # Col 3: 参数与方法 (Specifics 2) - 包含公式和图片
        p2 = row_cells[2].paragraphs[0]
        
        # 添加参数
        params_run = p2.add_run("1. 详细参数：\n")
        params_run.bold = True
        params_run.font.size = Pt(9)
        params_content = p2.add_run(f"{data.get('params', 'N/A')}")
        params_content.font.size = Pt(9)
        
        # 添加公式
        if 'formulas' in data and data['formulas']:
            formula_run = p2.add_run("\n\n2. 核心公式：\n")
            formula_run.bold = True
            formula_run.font.size = Pt(9)
            
            # 处理公式字符串，尝试渲染为图片
            formulas = data['formulas']
            if isinstance(formulas, str):
                formulas = [formulas]
                
            for formula in formulas:
                # 尝试渲染LaTeX为图片
                formula_img = self._render_latex_to_image(formula)
                
                if formula_img:
                    # 如果成功渲染，添加图片
                    formula_para = row_cells[2].add_paragraph()
                    formula_run = formula_para.add_run()
                    try:
                        formula_run.add_picture(formula_img, width=Inches(4))
                    except:
                        # 如果添加图片失败，显示LaTeX源码
                        formula_para.add_run(formula)
                else:
                    # 如果渲染失败，显示LaTeX源码
                    formula_code = p2.add_run(f"\n{formula}\n")
                    formula_code.font.size = Pt(8)
                    formula_code.font.name = 'Courier New'
        
        # 插入图片
        if image_stream:
            img_run = p2.add_run("\n\n3. 关键图表：\n")
            img_run.bold = True
            img_run.font.size = Pt(9)
            
            # 调整图片宽度以适应单元格
            img_para = row_cells[2].add_paragraph()
            img_run = img_para.add_run()
            try:
                img_run.add_picture(image_stream, width=Inches(3.5))
            except Exception as e:
                error_run = img_para.add_run(f"[图片加载失败: {str(e)}]")
                error_run.font.color.rgb = RGBColor(255, 0, 0)
        
        # Col 4: Comments
        p3 = row_cells[3].paragraphs[0]
        comments_run = p3.add_run("Comments:\n")
        comments_run.bold = True
        comments_run.font.size = Pt(9)
        comments_content = p3.add_run(data.get('comments', 'N/A'))
        comments_content.font.size = Pt(9)
        
        # Col 5: Why (Tags)
        p4 = row_cells[4].paragraphs[0]
        why_run = p4.add_run("Why:\n")
        why_run.bold = True
        why_run.font.size = Pt(9)
        tags_content = p4.add_run(data.get('why', data.get('tags', 'N/A')))
        tags_content.font.size = Pt(9)
        
        # 添加分页符（可选）
        # self.document.add_page_break()

    def add_summary_section(self, summary_data):
        """
        添加总结部分
        """
        # 添加总结标题
        self.document.add_heading('总结与分析', level=1)
        
        # 添加总结内容
        summary_para = self.document.add_paragraph()
        summary_run = summary_para.add_run(summary_data)
        summary_run.font.size = Pt(11)

    def save(self, filename="SPE_Analysis_Report.docx"):
        """
        保存Word文档
        
        Args:
            filename: 文件名
            
        Returns:
            str: 保存的文件名
        """
        self.document.save(filename)
        return filename
    
    def save_to_bytes(self):
        """
        将Word文档保存到字节流
        
        Returns:
            BytesIO: 包含Word文档的字节流
        """
        buffer = BytesIO()
        self.document.save(buffer)
        buffer.seek(0)
        return buffer